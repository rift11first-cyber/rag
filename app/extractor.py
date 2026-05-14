from __future__ import annotations

import csv
import io
import json
import shutil
from pathlib import Path

import openpyxl
import pytesseract
from docx import Document
from PIL import Image
from pypdf import PdfReader
import pypdfium2 as pdfium

from app.config import settings


class ExtractionError(RuntimeError):
    pass


def extract_text(file_name: str, payload: bytes, content_type: str | None = None) -> str:
    suffix = Path(file_name).suffix.lower()

    if settings.tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = settings.tesseract_cmd

    if suffix in {".txt", ".md", ".py", ".js", ".json", ".csv"}:
        return _extract_textish(suffix, payload)
    if suffix == ".docx":
        return _extract_docx(payload)
    if suffix in {".xlsx", ".xlsm"}:
        return _extract_xlsx(payload)
    if suffix == ".pdf":
        return _extract_pdf(payload)
    if suffix in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}:
        return _extract_image(payload)

    if content_type and content_type.startswith("text/"):
        return payload.decode("utf-8", errors="ignore")

    return payload.decode("utf-8", errors="ignore")


def _extract_textish(suffix: str, payload: bytes) -> str:
    if suffix == ".csv":
        rows = []
        decoded = payload.decode("utf-8", errors="ignore")
        reader = csv.reader(io.StringIO(decoded))
        for row in reader:
            rows.append(", ".join(cell.strip() for cell in row if cell.strip()))
        return "\n".join(row for row in rows if row)

    if suffix == ".json":
        parsed = json.loads(payload.decode("utf-8", errors="ignore"))
        return json.dumps(parsed, indent=2, ensure_ascii=True)

    return payload.decode("utf-8", errors="ignore")


def _extract_docx(payload: bytes) -> str:
    document = Document(io.BytesIO(payload))
    return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())


def _extract_xlsx(payload: bytes) -> str:
    workbook = openpyxl.load_workbook(io.BytesIO(payload), data_only=True)
    blocks: list[str] = []
    for sheet in workbook.worksheets:
        rows: list[str] = []
        for row in sheet.iter_rows(values_only=True):
            values = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
            if values:
                rows.append(" | ".join(values))
        if rows:
            blocks.append(f"{sheet.title}\n" + "\n".join(rows))
    return "\n\n".join(blocks)


def _extract_pdf(payload: bytes) -> str:
    reader = PdfReader(io.BytesIO(payload))
    page_text = [(page.extract_text() or "").strip() for page in reader.pages]
    merged = "\n\n".join(item for item in page_text if item)
    avg_chars = len(merged) / max(len(reader.pages), 1)
    if avg_chars >= 80:
        return merged
    return _ocr_pdf(payload)


def _extract_image(payload: bytes) -> str:
    _ensure_tesseract()
    image = Image.open(io.BytesIO(payload))
    return pytesseract.image_to_string(image).strip()


def _ocr_pdf(payload: bytes) -> str:
    _ensure_tesseract()
    document = pdfium.PdfDocument(payload)
    pages: list[str] = []
    for index in range(len(document)):
        page = document[index]
        bitmap = page.render(scale=2.0).to_pil()
        text = pytesseract.image_to_string(bitmap).strip()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def _ensure_tesseract() -> None:
    binary = settings.tesseract_cmd or shutil.which("tesseract")
    if not binary:
        raise ExtractionError("Tesseract OCR is required for scanned PDFs and images but `tesseract` was not found on PATH.")
